from __future__ import annotations
from datetime import datetime
from typing import Any
from pypdf import PdfReader
from docx import Document
from io import BytesIO
import tempfile
import re
import os
# import textract

import subprocess

def _run_cmd(cmd: list[str], timeout: int = 10) -> str:
    res = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout)
    if res.returncode != 0:
        raise RuntimeError(f"cmd failed: {' '.join(cmd)}\n{res.stderr}")
    return res.stdout

def extract_pdf_meta(data: bytes) -> dict[str, Any]:
    from io import BytesIO
    reader = PdfReader(BytesIO(data))
    info = reader.metadata or {}
    return {
        "pages": len(reader.pages),
        "author": str(info.get("/Author") or ""),
        "title": str(info.get("/Title") or ""),
        "creation_date": str(info.get("/CreationDate") or ""),
        "creator": str(info.get("/Creator") or ""),
        "producer": str(info.get("/Producer") or ""),
    }

def extract_docx_meta(data: bytes) -> dict[str, Any]:
    from io import BytesIO
    doc = Document(BytesIO(data))
    cp = doc.core_properties
    return {
        "paragraphs": sum(len(p.text) > 0 for p in doc.paragraphs),
        "tables": len(doc.tables),
        "title": cp.title or "",
        "author": cp.author or "",
        "created": str(cp.created or ""),
    }


def extract_doc_meta(data: bytes) -> dict[str, object]:
    """
    Best-effort для .doc без textract:
      - пытаемся через antiword (UTF-8)
      - при ошибке — через catdoc (-w)
      - считаем параграфы как блоки, таблицы — по строкам с >=2 табами
      - title/author/created могут быть пустыми (формат древний)
    """
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(data)
        path = tmp.name
    try:
        text = ""
        try:
            text = _run_cmd(["antiword", "-m", "UTF-8.txt", path])
        except Exception:
            text = _run_cmd(["catdoc", "-w", path])

        blocks = [b for b in re.split(r"\n\s*\n", text) if b.strip()]
        paragraphs = len(blocks)
        tables = sum(1 for line in text.splitlines() if line.count("\t") >= 2)

        return {
            "paragraphs": paragraphs,
            "tables": tables,
            "title": "",
            "author": "",
            "created": "",
        }
    finally:
        try:
            os.remove(path)
        except OSError:
            pass